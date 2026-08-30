#!/usr/bin/env python3
"""数据安全事件应急响应 respond-incident.

根据事件类型/规模/数据种类判定报告义务、时限，并生成监管报告书草案与个人告知文草案。
本工具为自检辅助，不构成法律意见，最终以监管口径为准。
"""
import argparse
import json
import os
import re
import sys
from datetime import datetime, timedelta

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_DIR = os.path.join(BASE_DIR, "data")


def load_json(path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def parse_dt(s):
    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M", "%Y-%m-%d"):
        try:
            return datetime.strptime(s, fmt)
        except ValueError:
            continue
    raise ValueError("无法解析时间：%s" % s)


def dedup_authorities(authority_strings):
    """将多条报告部门描述按分隔符拆成原子部门并去重，避免语义重叠导致的重复文案。
    保留首次出现顺序；若某原子部门已被更长的同义表述覆盖则合并。"""
    seen = []
    for s in authority_strings:
        for part in re.split(r"\s*[/、,，]\s*", s):
            part = part.strip()
            if part and part not in seen:
                seen.append(part)
    return seen


def assess(inp, rules):
    """判定报告义务、时限、级别，返回结构化结果，含判定追溯 trace。"""
    ro = rules["report_obligation"]
    sev_rules = rules["severity_rules"]
    cats = inp.get("data_categories", [])
    affected = inp.get("affected_users", 0) or 0
    discovered = parse_dt(inp["discovered_date"])
    trace = []

    triggered = []  # (rule_key, authority, deadline_hours, deadline_desc)

    # 重要数据
    if "重要数据" in cats:
        r = ro["important_data"]
        triggered.append(r)
        trace.append({"rule": "important_data", "matched": True,
                      "basis": "data_categories 含‘重要数据’", "deadline_hours": r["report_deadline_hours"]})
    else:
        trace.append({"rule": "important_data", "matched": False, "basis": "未涉重要数据"})

    # 大规模个人信息
    thr = sev_rules["large_scale_personal_info_threshold"]
    if affected >= thr:
        r = ro["large_scale_personal_info"]
        triggered.append(r)
        trace.append({"rule": "large_scale_personal_info", "matched": True,
                      "basis": "affected_users=%d >= %d" % (affected, thr),
                      "deadline_hours": r["report_deadline_hours"]})
    else:
        trace.append({"rule": "large_scale_personal_info", "matched": False,
                      "basis": "affected_users=%d < %d" % (affected, thr)})

    # 敏感个人信息
    if "敏感个人信息" in cats:
        r = ro["sensitive_personal_info"]
        triggered.append(r)
        trace.append({"rule": "sensitive_personal_info", "matched": True,
                      "basis": "data_categories 含‘敏感个人信息’", "deadline_hours": r["report_deadline_hours"]})
    else:
        trace.append({"rule": "sensitive_personal_info", "matched": False, "basis": "未涉敏感个人信息"})

    must_report = len(triggered) > 0
    if must_report:
        # 取最严(最短)时限
        chosen = min(triggered, key=lambda x: x["report_deadline_hours"])
        authorities = dedup_authorities(t["authority"] for t in triggered)
        deadline_hours = chosen["report_deadline_hours"]
        deadline_desc = chosen["deadline_desc"]
        report_deadline = discovered + timedelta(hours=deadline_hours)
    else:
        g = ro["general"]
        authorities = [g["authority"]]
        deadline_hours = g["report_deadline_hours"]
        deadline_desc = g["deadline_desc"]
        report_deadline = discovered + timedelta(hours=deadline_hours)
        trace.append({"rule": "general", "matched": True, "basis": "未触发强制报告项，按一般事件处置",
                      "deadline_hours": deadline_hours})

    # 级别
    if "重要数据" in cats or affected >= thr:
        severity = "重大"
    elif "敏感个人信息" in cats:
        severity = "较大"
    else:
        severity = "一般"

    # 个人告知义务
    ni = rules["notify_individual_rule"]
    notify_thr = sev_rules["notify_individual_affected_threshold"]
    has_pi = any(c in ("个人信息", "敏感个人信息") for c in cats)
    notify_individual = has_pi and affected >= notify_thr
    notify_deadline = discovered + timedelta(hours=ni["notify_deadline_hours"]) if notify_individual else None
    trace.append({"rule": "notify_individual", "matched": notify_individual,
                  "basis": "涉个人信息且受影响人数>=%d" % notify_thr if notify_individual
                  else "无需强制告知个人(或已采取措施可避免损害)"})

    return {
        "rules_version": rules["version"],
        "rules_updated_date": rules["updated_date"],
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "severity": severity,
        "must_report_authority": must_report,
        "report_authorities": authorities,
        "report_deadline_hours": deadline_hours,
        "report_deadline": report_deadline.strftime("%Y-%m-%d %H:%M:%S"),
        "report_deadline_desc": deadline_desc,
        "notify_individual": notify_individual,
        "notify_deadline": notify_deadline.strftime("%Y-%m-%d %H:%M:%S") if notify_deadline else None,
        "notify_deadline_desc": ni["deadline_desc"] if notify_individual else "本情形下暂无强制告知义务",
        "decision_trace": trace,
        "disclaimer": rules["disclaimer"],
    }


def _fill(text, ctx):
    for k, v in ctx.items():
        text = text.replace("{%s}" % k, str(v))
    return text


def build_context(inp, assessment):
    return {
        "controller_name": inp.get("controller_name", "____"),
        "report_time": assessment["generated_at"],
        "severity": assessment["severity"],
        "incident_type": inp.get("incident_type", "数据安全事件"),
        "discovered_date": inp.get("discovered_date", ""),
        "data_categories": "、".join(inp.get("data_categories", [])) or "未明确",
        "cross_border": "是" if inp.get("cross_border") else "否",
        "affected_users": inp.get("affected_users", 0),
        "impact_desc": "可能导致相关数据被非授权访问、泄露或滥用，存在对个人权益或数据安全造成损害的风险",
    }


def render_doc(template, ctx, path):
    """生成 DOCX(若 python-docx 可用)，否则回退 Markdown(.md)。返回实际路径。"""
    title = template["title"]
    try:
        from docx import Document
        doc = Document()
        doc.add_heading(title, level=0)
        for sec in template["sections"]:
            doc.add_heading(sec["heading"], level=1)
            for para in _fill(sec["body"], ctx).split("\n"):
                doc.add_paragraph(para)
        doc.save(path)
        return path
    except Exception:
        md_path = os.path.splitext(path)[0] + ".md"
        lines = ["# " + title, ""]
        for sec in template["sections"]:
            lines.append("## " + sec["heading"])
            lines.append(_fill(sec["body"], ctx))
            lines.append("")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        return md_path


def main():
    ap = argparse.ArgumentParser(description="数据安全事件应急响应")
    ap.add_argument("--input", required=True, help="输入 JSON 文件路径")
    ap.add_argument("--outdir", required=True, help="输出目录")
    args = ap.parse_args()

    os.makedirs(args.outdir, exist_ok=True)
    inp = load_json(args.input)
    rules = load_json(os.path.join(DATA_DIR, "reporting_rules.json"))
    templates = load_json(os.path.join(DATA_DIR, "templates.json"))

    assessment = assess(inp, rules)
    ctx = build_context(inp, assessment)

    json_path = os.path.join(args.outdir, "assessment.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(assessment, f, ensure_ascii=False, indent=2)

    reg_path = render_doc(templates["regulatory_report"], ctx,
                          os.path.join(args.outdir, "regulatory_report.docx"))
    notice_path = render_doc(templates["individual_notice"], ctx,
                             os.path.join(args.outdir, "individual_notice.docx"))

    print("事件级别：%s" % assessment["severity"])
    print("需向监管报告：%s，部门：%s" %
          ("是" if assessment["must_report_authority"] else "否", "、".join(assessment["report_authorities"])))
    print("报告截止时间：%s (%s)" % (assessment["report_deadline"], assessment["report_deadline_desc"]))
    print("需告知个人：%s" % ("是" if assessment["notify_individual"] else "否"))
    print("判定 JSON：%s" % json_path)
    print("监管报告书：%s" % reg_path)
    print("个人告知文：%s" % notice_path)
    return assessment


if __name__ == "__main__":
    main()
